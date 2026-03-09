import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Add Title
    title = doc.add_heading('Examen Final - Estrategias de Aprendizaje', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Instructions
    doc.add_paragraph('Nombre del alumno: __________________________________________________')
    doc.add_paragraph('Fecha: _____________________  Calificación: _________ / 10\n')
    
    doc.add_paragraph('Instrucciones: Lee cuidadosamente cada pregunta y subraya o encierra en un círculo la respuesta correcta.\n')

    quiz_data = [
        # Tema 1.1
        {
            "question": "¿Qué pilar de la educación consiste en adquirir destrezas y habilidades para trabajar en equipo (por ejemplo, saber aspirar secreciones)?",
            "options": [
                "a) Aprender a Conocer",
                "b) Aprender a Hacer",
                "c) Aprender a Vivir Juntos",
                "d) Aprender a Ser"
            ]
        },
        {
            "question": "¿Qué pilar de la educación se enfoca en la ética, la autonomía y la personalidad del enfermero?",
            "options": [
                "a) Aprender a Conocer",
                "b) Aprender a Hacer",
                "c) Aprender a Ser",
                "d) Aprender a Vivir Juntos"
            ]
        },
        # Tema 1.2
        {
            "question": "¿Cuál es el método de estudio que usa un reloj para estudiar por 25 minutos y descansar 5 minutos?",
            "options": [
                "a) Método Cornell",
                "b) Técnica Pomodoro",
                "c) Mapa Mental",
                "d) Sistema Kanban"
            ]
        },
        {
            "question": "¿Qué método para tomar apuntes divide la hoja en 3 partes: 'Pistas', 'Notas' y 'Resumen'?",
            "options": [
                "a) Sistema Kanban",
                "b) Método Pomodoro",
                "c) Método Cornell",
                "d) Esquema de llaves"
            ]
        },
        # Tema 1.3 / 1.4
        {
            "question": "Para formar un hábito de estudio fuerte, ¿cuáles son los 3 pasos de nuestro cerebro (Bucle Neurológico)?",
            "options": [
                "a) Motivación, Acción y Resultado",
                "b) Planear, Hacer y Revisar",
                "c) Señal, Rutina y Recompensa",
                "d) Leer, Escribir y Repasar"
            ]
        },
        {
            "question": "Según la Matriz del tiempo, ¿qué tipo de actividad es estudiar para el Postécnico, sabiendo que no es para mañana pero nos ayuda a crecer?",
            "options": [
                "a) Urgente e Importante",
                "b) No Urgente pero Importante",
                "c) Urgente pero No importante",
                "d) No Urgente y No Importante"
            ]
        },
        # Tema 1.5 / 1.6
        {
            "question": "En la Andragogía (cómo aprendemos los adultos), ¿cuál es nuestra mayor ventaja al estudiar temas nuevos?",
            "options": [
                "a) Que tenemos excelente memoria.",
                "b) Que tenemos mucha experiencia clínica previa de nuestro trabajo.",
                "c) Que nos gusta mucho leer libros largos.",
                "d) Que solo aprendemos si nos dan un premio."
            ]
        },
        {
            "question": "¿Qué tipo de luz se recomienda para tu lugar de estudio si no te quieres quedar dormido?",
            "options": [
                "a) Luz amarilla oscura",
                "b) Lámparas de colores",
                "c) Luz blanca/fría",
                "d) Poca luz, casi a oscuras"
            ]
        },
        # Tema 2.1
        {
            "question": "¿Cómo se llama el esquema que tiene un dibujo en el centro y coloridas ramas hacia los lados?",
            "options": [
                "a) Diagrama de Flujo",
                "b) Mapa Conceptual",
                "c) Mapa Mental",
                "d) Línea del tiempo"
            ]
        },
        {
            "question": "¿Qué tipo de lectura utilizamos si queremos encontrar si un artículo científico tiene errores o información falsa?",
            "options": [
                "a) Lectura Rápida",
                "b) Lectura de Comprensión",
                "c) Lectura Crítica",
                "d) Lectura Literal"
            ]
        },
        # Tema 2.2
        {
            "question": "Al usar la nemotecnia (acrónimo) V.O.M.I.T. para atender pacientes graves, ¿qué significa la 'O' y la 'M'?",
            "options": [
                "a) Orina y Medicamentos",
                "b) Oxígeno y Monitor EKG",
                "c) Observación y Movimiento",
                "d) Oídos y Manos"
            ]
        },
        {
            "question": "Si eres una persona auditiva, ¿qué técnica te sirve para estudiar mientras manejas hacia el hospital?",
            "options": [
                "a) Hacer mapas mentales",
                "b) Leer libros en el carro",
                "c) Escuchar notas de voz grabadas por ti mismo",
                "d) Ver videos musicales"
            ]
        },
        # Tema 2.3
        {
            "question": "¿Por qué es tan bueno practicar un 'Código Rojo' en un maniquí de simulación?",
            "options": [
                "a) Para jugar un rato.",
                "b) Para platicar con los compañeros.",
                "c) Para cansar a los alumnos.",
                "d) Para generar memoria en nuestros músculos y practicar bajo estrés leve."
            ]
        },
        {
            "question": "¿Cómo se llama la dinámica grupal donde todos dicen muchas ideas rápidas (sin burlarse) para resolver un problema?",
            "options": [
                "a) Lluvia de Ideas",
                "b) Debate",
                "c) Entrevista",
                "d) Lectura en voz alta"
            ]
        },
        # Tema 3.1 a 3.5
        {
            "question": "En el método S.Q.A. (Sé, Quiero saber, Aprendí), ¿Para qué sirve la columna de la 'A' (Aprendí)?",
            "options": [
                "a) Para apuntar las dudas iniciales.",
                "b) Para evaluar al final de la clase qué cosas nuevas logramos entender.",
                "c) Para escribir todo el apunte del maestro.",
                "d) Para calificar al maestro."
            ]
        },
        {
            "question": "¿Qué herramienta usarías para ver claramente las diferencias y semejanzas entre el paracetamol y el ibuprofeno?",
            "options": [
                "a) Una lluvia de ideas",
                "b) Un cuadro clínico",
                "c) Un cuadro comparativo",
                "d) Una infografía"
            ]
        },
        {
            "question": "En un Mapa Conceptual (que va de arriba hacia abajo), ¿qué elemento es obligatorio que conecte las ideas?",
            "options": [
                "a) Solo flechas sin nada escrito.",
                "b) Palabras de enlace (como: \"se divide en\", \"causa\").",
                "c) Muchos colores.",
                "d) Imágenes grandes."
            ]
        },
        # Tema 4.1 a 4.3
        {
            "question": "Para nunca perder tus tareas médicas aunque se rompa tu computadora o te roben la mochila, ¿dónde deberías guardarlas?",
            "options": [
                "a) En una memoria USB.",
                "b) Imprimir todo.",
                "c) Mandarlas por fax.",
                "d) En la Nube (Google Drive, OneDrive)."
            ]
        },
        {
            "question": "¿Qué papelillos cuadrados en blanco y negro podemos pegar al carro rojo para abrirlos con la cámara del celular y ver un video?",
            "options": [
                "a) Código de barras",
                "b) Stickers decorativos",
                "c) Códigos QR",
                "d) Etiquetas de precios"
            ]
        },
        {
            "question": "Hablando de TICS, ¿Qué significa la sigla M-learning?",
            "options": [
                "a) Aprender mucha Maternidad.",
                "b) Aprender Mecánica de enfermería.",
                "c) Mobile Learning: Aprender directamente desde el Teléfono Móvil (Celular).",
                "d) Aprender usando Maniquíes."
            ]
        }
    ]

    for i, q in enumerate(quiz_data, 1):
        # Adding the question text
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {q['question']}")
        run.bold = True
        
        # Adding options
        for opt in q['options']:
            p_opt = doc.add_paragraph(opt)
            p_opt.paragraph_format.left_indent = Pt(20)
            
        doc.add_paragraph() # spacing

    # Save document
    output_path = 'Examen_Final_Imprimible.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == '__main__':
    main()
