import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Add Title
    title = doc.add_heading('Resumen de Temas y Subtemas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Documento generado automáticamente a partir del contenido del curso.')
    
    with open('index.html', 'r', encoding='utf-8') as f:
        html = f.read()
        
    soup = BeautifulSoup(html, 'html.parser')
    
    sections = soup.find_all('section', class_='content-section')
    
    for section in sections:
        # Get section title
        h1 = section.find('h1')
        if not h1:
            continue
            
        heading = doc.add_heading(h1.get_text(strip=True), level=1)
        
        # Get all cards in this section
        cards = section.find_all('div', class_='card')
        for card in cards:
            # Subtopic title
            h3 = card.find(['h3', 'h4'])
            if h3:
                # Remove icon text if any, but get_text does fine usually. Let's just use get_text
                doc.add_heading(h3.get_text(strip=True), level=2)
                
            # Content inside card
            # Concepts are usually in <p> tags with <strong>
            paragraphs = card.find_all('p')
            for p in paragraphs:
                text = p.get_text(strip=True)
                if text:
                    p_element = doc.add_paragraph()
                    # Let's try to keep the bold formatting if it starts with Concepto, Detalle, etc.
                    strong = p.find('strong')
                    if strong and text.startswith(strong.get_text(strip=True)):
                        run = p_element.add_run(strong.get_text(strip=True) + " ")
                        run.bold = True
                        p_element.add_run(text[len(strong.get_text(strip=True)):].strip())
                    else:
                        p_element.add_run(text)
            
            # Lists (like details)
            ul = card.find('ul', class_='check-list')
            if ul:
                for li in ul.find_all('li'):
                    text = li.get_text(strip=True)
                    strong = li.find('strong')
                    p_element = doc.add_paragraph(style='List Bullet')
                    if strong and text.startswith(strong.get_text(strip=True)):
                        run = p_element.add_run(strong.get_text(strip=True) + " ")
                        run.bold = True
                        p_element.add_run(text[len(strong.get_text(strip=True)):].strip())
                    else:
                        p_element.add_run(text)
                        
            # Example box
            example = card.find('div', class_='example-box')
            if example:
                ex_p = doc.add_paragraph()
                run = ex_p.add_run("Ejemplo Práctico:")
                run.bold = True
                run.font.color.rgb = RGBColor(43, 108, 176) # A blue color
                
                # Exclude the <strong> text "Ejemplo Práctico:" from the rest of the text
                ex_text = example.get_text(separator=" ", strip=True)
                # Just appending the whole text is fine, but it might duplicate "Ejemplo Práctico".
                # Let's cleanly extract just the span or content.
                ex_span = example.find('span')
                if ex_span:
                    ex_p.add_run(" " + ex_span.get_text(strip=True))
                else:
                    # just clean it up 
                    clean_text = ex_text.replace("Ejemplo Práctico:", "").strip()
                    ex_p.add_run(" " + clean_text)
                    
            # Add a small spacing
            doc.add_paragraph()

    # Save document
    output_path = 'Resumen_Temas_Tecnologia.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == '__main__':
    main()
